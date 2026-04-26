from __future__ import annotations

from collections import defaultdict
from pathlib import Path
from zipfile import ZipFile
import xml.etree.ElementTree as ET

from docx import Document

from app.models.domain import DocumentPage, DocumentParseResult, ExtractedWord
from app.utils.text import build_sentence_occurrences, clean_marked_term, collapse_whitespace, has_english_content


XML_NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
}


def parse_docx(path: Path) -> DocumentParseResult:
    document = Document(path)
    paragraphs = [collapse_whitespace(paragraph.text) for paragraph in document.paragraphs if paragraph.text.strip()]
    full_text = "\n".join(paragraphs)
    words: list[ExtractedWord] = []
    for paragraph in document.paragraphs:
        words.extend(
            ExtractedWord(word=token, source="docx:red_text")
            for token in _collect_styled_phrases(paragraph.runs, _is_red_run)
        )
        words.extend(
            ExtractedWord(word=token, source="docx:underline")
            for token in _collect_styled_phrases(paragraph.runs, lambda run: bool(run.underline))
        )

    comment_words = _extract_commented_words(path)
    words.extend(ExtractedWord(word=token, source="docx:comment") for token in comment_words)
    deduped = _dedupe_extracted(words)
    pages = [DocumentPage(text=full_text, page_number=None)] if full_text else []
    return DocumentParseResult(
        source_type="docx",
        full_text=full_text,
        pages=pages,
        words=deduped,
        sentences=build_sentence_occurrences(pages),
    )


def _is_red_run(run) -> bool:
    color = getattr(run.font.color, "rgb", None)
    if color is None:
        return False
    return str(color).upper() in {"FF0000", "C00000", "FF3B30"}


def _collect_styled_phrases(runs, predicate) -> list[str]:
    phrases: list[str] = []
    buffer: list[str] = []
    for run in runs:
        text = run.text or ""
        if predicate(run):
            buffer.append(text)
            continue

        if buffer and text.isspace():
            buffer.append(text)
            continue

        phrase = clean_marked_term("".join(buffer))
        if phrase and has_english_content(phrase):
            phrases.append(phrase)
        buffer = []

    phrase = clean_marked_term("".join(buffer))
    if phrase and has_english_content(phrase):
        phrases.append(phrase)
    return phrases


def _extract_commented_words(path: Path) -> list[str]:
    try:
        with ZipFile(path) as archive:
            document_xml = archive.read("word/document.xml")
    except KeyError:
        return []
    root = ET.fromstring(document_xml)
    active: dict[str, list[str]] = defaultdict(list)
    collected: dict[str, list[str]] = defaultdict(list)
    for element in root.iter():
        tag = element.tag
        if tag.endswith("commentRangeStart"):
            comment_id = element.attrib.get(f"{{{XML_NS['w']}}}id")
            if comment_id is not None:
                active[comment_id]
        elif tag.endswith("commentRangeEnd"):
            comment_id = element.attrib.get(f"{{{XML_NS['w']}}}id")
            if comment_id is not None and comment_id in active:
                collected[comment_id].extend(active.pop(comment_id))
        elif tag.endswith("t"):
            text = element.text or ""
            for comment_id in list(active.keys()):
                active[comment_id].append(text)
    merged_texts = ["".join(parts) for parts in collected.values()]
    return [
        phrase
        for phrase in (clean_marked_term(text) for text in merged_texts)
        if phrase and has_english_content(phrase)
    ]


def _dedupe_extracted(words: list[ExtractedWord]) -> list[ExtractedWord]:
    seen: dict[str, ExtractedWord] = {}
    for item in words:
        key = item.word.strip().lower()
        if key and key not in seen:
            seen[key] = item
    return list(seen.values())
