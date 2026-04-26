from __future__ import annotations

import shutil
import uuid
from pathlib import Path

from docx import Document
from docx.shared import RGBColor

from app.services.document.docx import parse_docx


def _make_test_root() -> Path:
    root = Path(".runtime/test-docx-parser") / uuid.uuid4().hex
    root.mkdir(parents=True, exist_ok=True)
    return root


def test_parse_docx_preserves_contiguous_red_and_underlined_phrases() -> None:
    root = _make_test_root()
    docx_path = root / "lesson.docx"

    document = Document()
    paragraph = document.add_paragraph()
    red_one = paragraph.add_run("Growth ")
    red_one.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    red_two = paragraph.add_run("mindset")
    red_two.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    paragraph.add_run(" helps.")

    paragraph = document.add_paragraph()
    underline_one = paragraph.add_run("room ")
    underline_one.underline = True
    underline_two = paragraph.add_run("key")
    underline_two.underline = True
    paragraph.add_run(" is ready.")

    document.save(docx_path)

    result = parse_docx(docx_path)

    try:
        assert [item.word for item in result.words] == ["Growth mindset", "room key"]
        assert [item.text for item in result.sentences] == ["Growth mindset helps.", "room key is ready."]
        assert [item.page_number for item in result.sentences] == [None, None]
    finally:
        shutil.rmtree(root, ignore_errors=True)
